"""DOCX document parser using python-docx."""

import io
from docx import Document


class DocxParser:
    """Parser for DOCX documents using python-docx."""

    def parse(self, file_bytes: bytes) -> str:
        """
        Extract text content from DOCX bytes.

        Args:
            file_bytes: Raw DOCX file bytes

        Returns:
            Extracted text content as string
        """
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)

    def supports(self, filename: str) -> bool:
        """Check if this parser supports DOCX files."""
        return filename.lower().endswith(".docx")
